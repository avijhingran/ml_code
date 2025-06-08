import os
from docx import Document

def save_reports(results, output_path):
    os.makedirs(output_path, exist_ok=True)
    for result in results:
        doc = Document()
        doc.add_heading(f"Program Analysis: {result['filename']}", 0)

        doc.add_heading("1. Summary", level=1)
        doc.add_paragraph(result["summary"])

        doc.add_heading("2. Variables", level=1)
        doc.add_paragraph(result["variables"])

        doc.add_heading("3. Calls", level=1)
        doc.add_paragraph(result["calls"])

        doc.add_heading("4. Mermaid Flowchart", level=1)
        doc.add_paragraph(result["mermaid"])

        save_path = os.path.join(output_path, f"{result['filename']}.docx")
        doc.save(save_path)
